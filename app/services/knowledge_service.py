"""Knowledge base ingestion and retrieval service."""

import csv
import io
import re
import uuid
from typing import List

from sqlalchemy.orm import Session

from app.models import KnowledgeChunk, KnowledgeDocument


class KnowledgeService:
    ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "csv"}

    def validate_extension(self, filename: str) -> str:
        ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
        if ext not in self.ALLOWED_EXTENSIONS:
            raise ValueError("Unsupported file type. Allowed: pdf, docx, txt, csv")
        return ext

    def extract_text(self, filename: str, payload: bytes) -> str:
        ext = self.validate_extension(filename)
        if ext == "txt":
            return payload.decode("utf-8", errors="ignore")
        if ext == "csv":
            decoded = payload.decode("utf-8", errors="ignore")
            reader = csv.reader(io.StringIO(decoded))
            rows = [" | ".join(cell.strip() for cell in row if cell is not None) for row in reader]
            return "\n".join(row for row in rows if row.strip())
        if ext == "pdf":
            try:
                from pypdf import PdfReader  # type: ignore
            except Exception as exc:
                raise ValueError("PDF support requires 'pypdf' package") from exc

            reader = PdfReader(io.BytesIO(payload))
            texts = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(texts)
        if ext == "docx":
            try:
                from docx import Document  # type: ignore
            except Exception as exc:
                raise ValueError("DOCX support requires 'python-docx' package") from exc

            doc = Document(io.BytesIO(payload))
            return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
        return ""

    def chunk_text(self, text: str, chunk_size: int = 700, overlap: int = 120) -> List[str]:
        normalized = re.sub(r"\s+", " ", (text or "").strip())
        if not normalized:
            return []

        chunks: List[str] = []
        cursor = 0
        while cursor < len(normalized):
            end = min(len(normalized), cursor + chunk_size)
            chunk = normalized[cursor:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(normalized):
                break
            cursor = max(end - overlap, cursor + 1)
        return chunks

    def ingest_document(
        self,
        db: Session,
        user_id: str,
        organization_id: str,
        filename: str,
        text: str,
        title: str | None = None,
    ) -> KnowledgeDocument:
        file_type = self.validate_extension(filename)
        chunks = self.chunk_text(text)
        if not chunks:
            raise ValueError("Could not extract text from uploaded document")

        doc = KnowledgeDocument(
            id=str(uuid.uuid4()),
            user_id=user_id,
            organization_id=organization_id,
            filename=filename,
            file_type=file_type,
            title=(title or filename),
            total_chunks=len(chunks),
        )
        db.add(doc)

        for idx, chunk in enumerate(chunks):
            db.add(
                KnowledgeChunk(
                    id=str(uuid.uuid4()),
                    document_id=doc.id,
                    organization_id=organization_id,
                    chunk_index=idx,
                    content=chunk,
                    token_count=len(chunk.split()),
                )
            )

        return doc

    def search_chunks(self, db: Session, organization_id: str, query: str, top_k: int = 5) -> list[KnowledgeChunk]:
        tokens = [token for token in re.findall(r"\w+", (query or "").lower()) if len(token) > 2]
        if not tokens:
            return []

        candidates = (
            db.query(KnowledgeChunk)
            .filter(KnowledgeChunk.organization_id == organization_id)
            .order_by(KnowledgeChunk.created_at.desc())
            .limit(800)
            .all()
        )

        scored = []
        for chunk in candidates:
            haystack = (chunk.content or "").lower()
            score = sum(1 for token in tokens if token in haystack)
            if score > 0:
                scored.append((score, chunk))

        scored.sort(key=lambda item: item[0], reverse=True)
        return [chunk for _, chunk in scored[: max(1, min(top_k, 10))]]

    def synthesize_answer(self, query: str, chunks: list[KnowledgeChunk]) -> str:
        if not chunks:
            return "Nao encontrei informacao relevante nos documentos enviados para responder isso."

        excerpts = []
        for idx, chunk in enumerate(chunks, start=1):
            excerpt = chunk.content[:260].strip()
            excerpts.append(f"[{idx}] {excerpt}")

        refs = "\n".join(excerpts)
        return (
            f"Com base nos documentos enviados, esta e a melhor resposta para: '{query}'.\n\n"
            f"Trechos relevantes:\n{refs}"
        )

    def list_documents(self, db: Session, organization_id: str, limit: int = 100) -> list[KnowledgeDocument]:
        safe_limit = max(1, min(limit, 500))
        return (
            db.query(KnowledgeDocument)
            .filter(KnowledgeDocument.organization_id == organization_id)
            .order_by(KnowledgeDocument.created_at.desc())
            .limit(safe_limit)
            .all()
        )

    def delete_document(self, db: Session, organization_id: str, document_id: str) -> bool:
        row = (
            db.query(KnowledgeDocument)
            .filter(
                KnowledgeDocument.id == document_id,
                KnowledgeDocument.organization_id == organization_id,
            )
            .first()
        )
        if not row:
            return False

        db.query(KnowledgeChunk).filter(KnowledgeChunk.document_id == row.id).delete()
        db.delete(row)
        return True
